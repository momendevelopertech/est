from __future__ import annotations

import json
from pathlib import Path
import sys
import zipfile

import openpyxl
import olefile
import xlrd
from docx import Document


BASE = Path(r"e:\est files")


def normalize(value):
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def sample_rows_openpyxl(sheet, limit=12):
    rows = []
    for row in sheet.iter_rows(min_row=1, max_row=min(sheet.max_row, limit), values_only=True):
        values = [normalize(cell) for cell in row]
        if any(values):
            rows.append(values)
    return rows


def analyze_xlsx(path: Path):
    wb = openpyxl.load_workbook(path, data_only=True)
    result = {"file": path.name, "type": path.suffix.lower(), "sheets": []}
    for name in wb.sheetnames:
        sheet = wb[name]
        result["sheets"].append(
            {
                "name": name,
                "max_row": sheet.max_row,
                "max_col": sheet.max_column,
                "sample_rows": sample_rows_openpyxl(sheet),
            }
        )
    return result


def analyze_xls(path: Path):
    wb = xlrd.open_workbook(path)
    result = {"file": path.name, "type": path.suffix.lower(), "sheets": []}
    for idx in range(wb.nsheets):
        sheet = wb.sheet_by_index(idx)
        rows = []
        for rx in range(min(sheet.nrows, 12)):
            values = [normalize(sheet.cell_value(rx, cx)) for cx in range(sheet.ncols)]
            if any(values):
                rows.append(values)
        result["sheets"].append(
            {
                "name": sheet.name,
                "max_row": sheet.nrows,
                "max_col": sheet.ncols,
                "sample_rows": rows,
            }
        )
    return result


def analyze_docx(path: Path):
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows[:12]:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                rows.append(values)
        tables.append({"rows": rows})
    return {
        "file": path.name,
        "type": path.suffix.lower(),
        "paragraphs": paragraphs[:20],
        "tables": tables,
    }


def sniff_file_kind(path: Path):
    with path.open("rb") as handle:
        header = handle.read(8)
    if header.startswith(b"PK"):
        return "zip"
    if header.startswith(b"\xd0\xcf\x11\xe0"):
        return "ole"
    return "unknown"


def analyze_doc_ole(path: Path):
    result = {"file": path.name, "type": path.suffix.lower(), "storage": "ole", "streams": []}
    if not olefile.isOleFile(path):
        result["error"] = "not_ole"
        return result
    ole = olefile.OleFileIO(path)
    try:
        for entry in ole.listdir()[:20]:
            result["streams"].append("/".join(entry))
    finally:
        ole.close()
    return result


def analyze_path(path: Path):
    suffix = path.suffix.lower()
    try:
        if suffix in {".xlsx", ".xlsm"}:
            return analyze_xlsx(path)
        if suffix == ".xls":
            return analyze_xls(path)
        if suffix == ".docx":
            kind = sniff_file_kind(path)
            if kind == "zip":
                return analyze_docx(path)
            if kind == "ole":
                return analyze_doc_ole(path)
            return {
                "file": path.name,
                "type": suffix,
                "storage": kind,
                "error": "unsupported_doc_container",
            }
        return {"file": path.name, "type": suffix, "error": "unsupported_extension"}
    except zipfile.BadZipFile:
        return {
            "file": path.name,
            "type": suffix,
            "storage": "zip",
            "error": "bad_zip_file",
        }
    except Exception as exc:
        return {"file": path.name, "type": suffix, "error": type(exc).__name__, "message": str(exc)}


def main():
    analyses = []
    for path in sorted(BASE.iterdir()):
        analyses.append(analyze_path(path))
    payload = json.dumps(analyses, ensure_ascii=False, indent=2)
    sys.stdout.buffer.write(payload.encode("utf-8"))


if __name__ == "__main__":
    main()
